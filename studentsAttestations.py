import openpyxl
import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING

# Charger le fichier Excel
workbook = openpyxl.load_workbook('C:/Users/nelson.PHEOC/Desktop/projects/python/studentPassasionsCertificats/resources/excelFiles/LISTE DES ETUDIANTS ESMA-MZ 2023-2024 DRSPC.xlsx')
sheet = workbook.active

# Créer le document Word
document = docx.Document()





#START OF THE HEADER PART !!!!!!!!!!!!!!!!!
# Access the header of the first section
header = document.sections[0].header

#setting the header top margin
document.sections[0].header_distance = Pt(0)
document.sections[0].left_margin = Inches(1)
document.sections[0].right_margin = Inches(1)

# Create a table in the header
htable = header.add_table(1, 3, Inches(15))
htable.style.space_before = Pt(0)  # Set the space before the table to 0



htab_cells = htable.rows[0].cells

# Add the left-aligned structured text to the first cell
ht0 = htab_cells[0].add_paragraph()
ht0.paragraph_format.space_before = Pt(0)  # Remove extra space before paragraph
ht0.text = "REPUBLIQUE DU CAMEROUN"
ht0.style.font.size = Pt(8)
ht0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht0.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht0.paragraph_format.line_spacing = 1
ht0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht0 = htab_cells[0].add_paragraph()
ht0.text = "Paix – Travail – Patrie"
ht0.style.font.size = Pt(8)
ht0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht0.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht0.paragraph_format.line_spacing = 1
ht0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht0 = htab_cells[0].add_paragraph()
ht0.text = "-----"
ht0.style.font.size = Pt(8)
ht0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht0.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht0.paragraph_format.line_spacing = 1
ht0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht0 = htab_cells[0].add_paragraph()
ht0.text = "MINISTERE DE LA SANTE PUBLIQUE"
ht0.style.font.size = Pt(8)
ht0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht0.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht0.paragraph_format.line_spacing = 1
ht0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht0 = htab_cells[0].add_paragraph()
ht0.text = "-----"
ht0.style.font.size = Pt(8)
ht0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht0.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht0.paragraph_format.line_spacing = 1
ht0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht0 = htab_cells[0].add_paragraph()
ht0.text = "DIRECTION DES RESSOURCES HUMAINES"
ht0.style.font.size = Pt(8)
ht0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht0.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht0.paragraph_format.line_spacing = 1
ht0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht0 = htab_cells[0].add_paragraph()
ht0.text = "-----"
ht0.style.font.size = Pt(8)
ht0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht0.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht0.paragraph_format.line_spacing = 1
ht0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht0 = htab_cells[0].add_paragraph()
ht0.text = "ÉCOLE DES SCIENCES MEDICALES ET D'APPLICATION MARIE ZAMBO"
ht0.style.font.size = Pt(8)
ht0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht0.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht0.paragraph_format.line_spacing = 1
ht0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht0 = htab_cells[0].add_paragraph()
ht0.text = "BP :"
ht0.style.font.size = Pt(8)
ht0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht0.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht0.paragraph_format.line_spacing = 1
ht0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht0 = htab_cells[0].add_paragraph()
ht0.text = "Tel : +237 655 666 140"
ht0.style.font.size = Pt(8)
ht0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht0.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht0.paragraph_format.line_spacing = 1
ht0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

# Add the image to the second cell
kh = htab_cells[1].add_paragraph()
kh.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph
kh.add_run().add_picture('C:/Users/nelson.PHEOC/Desktop/projects/python/studentPassasionsCertificats/resources/esmaLogo.png', width=Inches(1.15))
kh.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add the RIGHT-aligned structured text to the third cell
ht1 = htab_cells[2].add_paragraph()
ht1.paragraph_format.space_before = Pt(0)  # Remove extra space before paragraph
ht1.text = "REPUBLIC OF CAMEROON"
ht1.style.font.size = Pt(8)
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht1.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht1.paragraph_format.line_spacing = 1
ht1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht1 = htab_cells[2].add_paragraph()
ht1.text = "Peace - Work - Fatherland"
ht1.style.font.size = Pt(8)
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht1.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht1.paragraph_format.line_spacing = 1
ht1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht1 = htab_cells[2].add_paragraph()
ht1.text = "-----"
ht1.style.font.size = Pt(8)
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht1.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht1.paragraph_format.line_spacing = 1
ht1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht1 = htab_cells[2].add_paragraph()
ht1.text = "MINISTRY OF PUBLIC HEALTH"
ht1.style.font.size = Pt(8)
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht1.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht1.paragraph_format.line_spacing = 1
ht1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht1 = htab_cells[2].add_paragraph()
ht1.text = "-----"
ht1.style.font.size = Pt(8)
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht1.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht1.paragraph_format.line_spacing = 1
ht1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht1 = htab_cells[2].add_paragraph()
ht1.text = "HUMAN RESOURCES DEPARTMENT"
ht1.style.font.size = Pt(8)
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht1.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht1.paragraph_format.line_spacing = 1
ht1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht1 = htab_cells[2].add_paragraph()
ht1.text = "-----"
ht1.style.font.size = Pt(8)
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht1.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht1.paragraph_format.line_spacing = 1
ht1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht1 = htab_cells[2].add_paragraph()
ht1.text = "SCHOOL OF MEDICAL AND APPLIED SCIENCES MARIE ZAMBO"
ht1.style.font.size = Pt(8)
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht1.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht1.paragraph_format.line_spacing = 1
ht1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht1 = htab_cells[2].add_paragraph()
ht1.text = "PO Box:"
ht1.style.font.size = Pt(8)
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht1.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht1.paragraph_format.line_spacing = 1
ht1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

ht1 = htab_cells[2].add_paragraph()
ht1.text = "Tel: +237 655 666 140"
ht1.style.font.size = Pt(8)
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ht1.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ht1.paragraph_format.line_spacing = 1
ht1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

#END OF THE HEADER PART !!!!!!!!!!!!!!!!!


# Load the Word document
# Find the last used row
last_used_row = 0
for row in range(1, sheet.max_row + 1):
    if any(sheet.cell(row=row, column=col).value for col in range(1, sheet.max_column + 1)):
        last_used_row = row

# Parcourir les lignes du fichier Excel
for row in range(2, last_used_row + 1):
    # Récupérer les informations de l'étudiant
    matricule = sheet.cell(row=row, column=2).value
    nom_prenom = sheet.cell(row=row, column=3).value
    sexe = sheet.cell(row=row, column=4).value
    date_naissance = sheet.cell(row=row, column=5).value.strftime('%m/%d/%Y') # Extract the date portion
    lieu_naissance = sheet.cell(row=row, column=6).value
    numero_telephone = sheet.cell(row=row, column=7).value
    statut = sheet.cell(row=row, column=8).value
    observations = sheet.cell(row=row, column=9).value
    annee = "2023/2024"
    filliere = "AID SOIGNANTS GENERALISTES"
    # Ajouter une nouvelle page au document Word
    title = document.add_heading(f"FICHE DE RENSEIGNEMENT DE CANDIDATURE A L'EXAMEN DE CERTFICATION ANNEE ACADEMIQUE 2023-2024", level=2)
    
    title.style.font.color.rgb = RGBColor(0,0,0)

    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    titleStylingRun = title.runs[0]
    
    titleStylingRun.bold = True
    titleStylingRun.underline = True
    
    
    p1 = document.add_paragraph(f"EXAMEN CONCERNE : {filliere}")
    p1.paragraph_format.space_before = Inches(0.1)
    p1Run = p1.runs[0]
    p1Run.bold = True
    p1Font = p1Run.font
    p1Font.name = 'Arial'
    p1Font.size = Pt(11)


    p2 = document.add_paragraph(f"IDENTIFICATION DU CANDIDAT")
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p2Run = p2.runs[0]
    p2Run.bold = True
    p2Font = p2Run.font
    p2Font.name = 'Arial'
    p2Font.size = Pt(11)
    
    p3 = document.add_paragraph("")
    p3.paragraph_format.left_indent = Inches(0.2)
    p3.paragraph_format.line_spacing = 1
    p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p3span1 = p3.add_run("■      ")
    p3span1.font.size = Pt(8)
    p3span2 = p3.add_run(f"Nom(s) et Prénom(s):   {nom_prenom}")
    p3span2.bold = True
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p3Run = p3.runs[0]
    p3Run.bold = True
    p3span2Font = p3span2.font
    p3span2Font.name = 'Arial'
    p3span2Font.size = Pt(11)
    
    p4 = document.add_paragraph("")
    p4.paragraph_format.left_indent = Inches(0.2)
    p4.paragraph_format.line_spacing = 1
    p4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p4span1 = p4.add_run("■      ")
    p4span1.font.size = Pt(8)
    p4span2 = p4.add_run(f"Né(e) le      {date_naissance}         à      {lieu_naissance}      Matricule: {matricule}")
    p4span2.bold = True
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p4Run = p4.runs[0]
    p4Run.bold = True
    p4span2Font = p4span2.font
    p4span2Font.name = 'Arial'
    p4span2Font.size = Pt(11)

    p5 = document.add_paragraph(f"ANNÉE D’ENTRÉE À L’ÉTABLISSEMENT : {annee}")
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p5Run = p5.runs[0]
    p5Run.bold = True
    p5Font = p5Run.font
    p5Font.name = 'Arial'
    p5Font.size = Pt(11)

    p6 = document.add_paragraph()
    p6.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # Add the regular text
    p6.add_run("SCOLARITE       Suffisante     ")

    # Add the special span
    span = p6.add_run("OUI")
    span.font.size = Pt(14)
    span.font.name = 'Arial'
    span.bold = True
    span.font.color.rgb = RGBColor(0, 0, 0)

    # Add borders to the second text element
    for i in range(len(p6.runs)):
        if i == 0.7:
            p6.runs[i].font.underline = True
            p6.runs[i].font.border = True
            p6.runs[i].font.border_size = Pt(2)
            p6.runs[i].font.border_color = (0, 0, 0)  # Set the border color to black
        
    p6Run = p6.runs[0]
    p6Run.bold = True
    p6Font = p6Run.font
    p6Font.name = 'Arial'
    p6Font.size = Pt(11)

    # table = document.add_table(1,2, Inches(5))
    # table.cell(0,0).text = 'Left Text'
    # table.cell(0,1).text = 'Right Text'

    # table.rows[0].cells[0].width = Pt(200)
    # table.rows[0].cells[1].width = Pt(50)

    # def set_cell_border(cell, **kwargs):
    #     """
    #     Set cell`s border
    #     Usage:

    #     set_cell_border(
    #         cell,
    #         top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
    #         bottom={"sz": 12, "color": "#00FF00", "val": "single"},
    #         start={"sz": 24, "val": "dashed", "shadow": "true"},
    #         end={"sz": 12, "val": "dashed"},
    #     )
    #     """
    #     tc = cell._tc
    #     tcPr = tc.get_or_add_tcPr()

    #     # check for tag existnace, if none found, then create one
    #     tcBorders = tcPr.first_child_found_in("w:tcBorders")
    #     if tcBorders is None:
    #         tcBorders = BaseOxmlElement('w:tcBorders')
    #         tcPr.append(tcBorders)

    #     # list over all available tags
    #     for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
    #         edge_data = kwargs.get(edge)
    #         if edge_data:
    #             tag = 'w:{}'.format(edge)

    #             # check for tag existnace, if none found, then create one
    #             element = tcBorders.find(qn(tag))
    #             if element is None:
    #                 element = BaseOxmlElement(tag)
    #                 tcBorders.append(element)

    #             # looks like order of attributes is important
    #             for key in ["sz", "val", "color", "space", "shadow"]:
    #                 if key in edge_data:
    #                     element.set(qn('w:{}'.format(key)), str(edge_data[key]))
    
    p7 = document.add_paragraph("")
    p7.paragraph_format.left_indent = Inches(0.2)
    p7.paragraph_format.line_spacing = 1
    p7.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p7span1 = p7.add_run("■      ")
    p7span1.font.size = Pt(8)
    p7span2 = p7.add_run(f"Nombre de jours de stage à récupérer :    RAS")
    p7span2.bold = True
    p7.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p7Run = p7.runs[0]
    p7Run.bold = True
    p7span2Font = p7span2.font
    p7span2Font.name = 'Arial'
    p7span2Font.size = Pt(11)
    
    p8 = document.add_paragraph("")
    p8.paragraph_format.left_indent = Inches(0.2)
    p8.paragraph_format.line_spacing = 1
    p8.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p8span1 = p8.add_run("■      ")
    p8span1.font.size = Pt(8)
    p8span2 = p8.add_run(f"Nombre de jours d’absence aux cours théoriques :  RAS")
    p8span2.bold = True
    p8.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p8Run = p8.runs[0]
    p8Run.bold = True
    p8span2Font = p8span2.font
    p8span2Font.name = 'Arial'
    p8span2Font.size = Pt(11)
    
    p9 = document.add_paragraph("")
    p9.paragraph_format.left_indent = Inches(0.2)
    p9.paragraph_format.line_spacing = 1
    p9.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p9span1 = p9.add_run("■      ")
    p9span1.font.size = Pt(8)
    p9span2 = p9.add_run(f"Nombre de gardes et permanences non effectuées :  RAS")
    p9span2.bold = True
    p9.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p9Run = p9.runs[0]
    p9Run.bold = True
    p9span2Font = p9span2.font
    p9span2Font.name = 'Arial'
    p9span2Font.size = Pt(11)
    
    p10 = document.add_paragraph("")
    p10.paragraph_format.left_indent = Inches(0.2)
    p10.paragraph_format.line_spacing = 1
    p10.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p10span1 = p10.add_run("■      ")
    p10span1.font.size = Pt(8)
    p10span2 = p10.add_run(f"Nombre de modules non validés :  {0}")
    p10span2.bold = True
    p10.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p10Run = p10.runs[0]
    p10Run.bold = True
    p10span2Font = p10span2.font
    p10span2Font.name = 'Arial'
    p10span2Font.size = Pt(11)
    
    p11 = document.add_paragraph("")
    p11.paragraph_format.left_indent = Inches(0.2)
    p11.paragraph_format.line_spacing = 1
    p11.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p11span1 = p11.add_run("■      ")
    p11span1.font.size = Pt(8)
    p11span2 = p11.add_run(f"Pourcentage de validation des modules :  100%")
    p11span2.bold = True
    p11.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p11Run = p11.runs[0]
    p11Run.bold = True
    p11span2Font = p11span2.font
    p11span2Font.name = 'Arial'
    p11span2Font.size = Pt(11)

    
    p12 = document.add_paragraph(f"DISCIPLINE")
    p12.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p12Run = p12.runs[0]
    p12Run.bold = True
    p12Font = p12Run.font
    p12Font.name = 'Arial'
    p12Font.size = Pt(11)
    
    p13 = document.add_paragraph("")
    p13.paragraph_format.left_indent = Inches(0.2)
    p13.paragraph_format.line_spacing = 1
    p13.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p13span1 = p13.add_run("■      ")
    p13span1.font.size = Pt(8)
    p13span2 = p13.add_run(f"Nombre de fois que l’élève a été traduit devant le conseil de discipline : {0}")
    p13span2.bold = True
    p13.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p13Run = p13.runs[0]
    p13Run.bold = True
    p13span2Font = p13span2.font
    p13span2Font.name = 'Arial'
    p13span2Font.size = Pt(11)
    
    p14 = document.add_paragraph("")
    p14.paragraph_format.left_indent = Inches(0.2)
    p14.paragraph_format.line_spacing = 1
    p14.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p14span1 = p14.add_run("■       ")
    p14span1.font.size = Pt(8)
    p14span2 = p14.add_run(f"Types de sanction(s) infligée(s) :   RAS")
    p14span2.bold = True
    p14.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p14Run = p14.runs[0]
    p14Run.bold = True
    p14span2Font = p14span2.font
    p14span2Font.name = 'Arial'
    p14span2Font.size = Pt(11)
    
    p15 = document.add_paragraph(f"FRAIS DE SCOLARITE")
    p15.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p15Run = p15.runs[0]
    p15Run.bold = True
    p15Font = p15Run.font
    p15Font.name = 'Arial'
    p15Font.size = Pt(11)

    # pSPECIAL = document.add_paragraph(f"try")
    
    # # Add a textbox to the paragraph
    # prun = pSPECIAL.add_run()
    # prun.add_picture('C:/Users/nelson.PHEOC/Desktop/projects/python/studentPassasionsCertificats.png', width=Inches(2), height=Inches(1))
    # prun.add_text

    # # Align the textbox
    # pSPECIAL.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # # Create a square bullet list
    # bullet_paragraph = document.add_paragraph()
    # bullet_paragraph.style.font.name = 'Wingdings'
    # bullet_paragraph.style.font.size = Pt(1
    #p13.paragraph_format.left_indent = Inches(0.2)
    # fisr = bullet_paragraph.add_run(" ■      ")
    # fisr.font.size = Pt(8)
    # bullet_list = bullet_paragraph.add_run("First item")
    

    
    p16 = document.add_paragraph("")
    p16.paragraph_format.left_indent = Inches(0.2)
    p16.paragraph_format.line_spacing = 1
    p16.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p16span1 = p16.add_run("■       ")
    p16span1.font.size = Pt(8)
    p16.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # Add the regular text
    p16span2 = p16.add_run("Totalement acquittés :       ")
    p16span2.bold = True
    p16span2Font = p16span2.font
    p16span2Font.name = 'Arial'
    p16span2Font.size = Pt(11)

    # Add the special span
    span1 = p16.add_run("OUI       ")
    span1.font.size = Pt(14)
    span1.font.name = 'Arial'
    span1.bold = True
    span1.font.color.rgb = RGBColor(0, 0, 0)

    span2 = p16.add_run('Somme restante :  0')
    span2.font.size = Pt(11)
    span2.font.name = 'Arial'
    span2.bold = True
    span2.font.color.rgb = RGBColor(0, 0, 0)
     
    
    p17 = document.add_paragraph("Autorisé(e) à composer :     ")
    p17.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p17Run = p17.runs[0]
    p17Run.bold = True
    p17Font = p17Run.font
    p17Font.name = 'Arial'
    p17Font.size = Pt(11)
    
    
    # Add the special span
    span1 = p17.add_run("OUI")
    span1.font.size = Pt(14)
    span1.font.name = 'Arial'
    span1.bold = True
    span1.font.color.rgb = RGBColor(0, 0, 0)
    
        
    p18 = document.add_paragraph("")
    p18.paragraph_format.left_indent = Inches(0.2)
    p18.paragraph_format.line_spacing = 1
    p18.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p18span1 = p18.add_run("■       ")
    p18span1.font.size = Pt(8)
    p18span2 = p18.add_run(f"OBSERVATIONS GENERALES : appliquée à ses études.")
    p18.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p18Run = p18.runs[0]
    p18span2.bold = True
    p18span2Font = p18span2.font
    p18span2Font.name = 'Arial'
    p18span2Font.size = Pt(11)
    
    
    p19 = document.add_paragraph(f"LA DIRECTRICE")
    p19.paragraph_format.space_before = Inches(0.2)
    p19.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p19Run = p19.runs[0]
    p19Run.bold = True
    p19Run.underline = True
    p19Font = p19Run.font
    p19Font.name = 'Arial'
    p19Font.size = Pt(10)

    p20 = document.add_paragraph(f"ELIMBI Madeleine Yvonne")
    p20.paragraph_format.space_before = Inches(0.5)
    p20.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p20Run = p20.runs[0]
    p20Run.bold = True
    p20Font = p20Run.font
    p20Font.name = 'Arial'
    p20Font.size = Pt(10)
    
    
    
    

    # Ajouter une nouvelle page
    document.add_page_break()


#START OF THE FOOTER PART !!!!!!!!!!!!!!!!!
# Access the footer of the first section
footer = document.sections[0].footer

# Create a table in the footer
ftable = footer.add_table(1, 2, Inches(6))


ftable.style.space_before = Pt(0)  # Set the space before the table to 0

ftable.rows[0].cells[0].width = Pt(50)
ftable.rows[0].cells[1].width = Pt(500)

ftab_cells = ftable.rows[0].cells

# Add the image to the first cell
kf = ftab_cells[0].add_paragraph()
kf.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph
kf.add_run().add_picture('C:/Users/nelson.PHEOC/Desktop/projects/python/studentPassasionsCertificats/resources/esmaLogo.png', width=Inches(0.3))
kf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


ft1 = ftab_cells[1].add_paragraph()
ft1.text = "ÉCOLE DES SCIENCES MEDICALES ET D’APPLICATION MARIE ZAMBO"
ft1.style.font.name = 'Arial'
ft1.style.font.size = Pt(8)
ft1.paragraph_format.line_spacing = 1
ft1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
ft1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ft1.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph

ft1 = ftab_cells[1].add_paragraph()
ft1.text = "Décision N°2531/D/MINSANTE/SG/DRH du 02 août 2021"
ft1.style.font.size = Pt(8)
ft1.style.font.name = 'Arial'
ft1.paragraph_format.line_spacing = 1
ft1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
ft1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
ft1.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraph
#END OF THE FOOTER PART !!!!!!!!!!!!!!!!!

# Enregistrer le document Word
document.save('C:/Users/nelson.PHEOC/Desktop/projects/python/studentPassasionsCertificats/results/FICHE_DE_RENSEIGNEMENT_DE_CANDIDATS.docx')